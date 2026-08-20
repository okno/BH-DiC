from __future__ import annotations

import io
import zipfile
from datetime import UTC, date, datetime
from decimal import Decimal

import pytest
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from bh_dic.exports import EmployeeExportRow, ExportFormat, HrExportService
from bh_dic.exports.service import ExportGenerationError, ascii_table, split_ascii_for_discord


def _rows(count: int = 3) -> tuple[EmployeeExportRow, ...]:
    return tuple(
        EmployeeExportRow(
            first_name=f"Nome {index}",
            last_name=f"Cognome {index:03d}",
            employee_id=f"EMP-SYNTH-{index:03d}",
            status="active" if index % 2 else "inactive",
            contract_expiry=date(2026, 12, min(28, index + 1)),
            contract_type="fixed_term · Full time",
            monthly_net=Decimal("1750.00") if index == 1 else None,
        )
        for index in range(1, count + 1)
    )


@pytest.mark.parametrize("export_format", list(ExportFormat))
def test_real_export_formats_are_generated_and_reopenable(export_format: ExportFormat) -> None:
    service = HrExportService(max_bytes=8 * 1024 * 1024)
    generated = service.generate_employees(
        _rows(65),
        export_format=export_format,
        created_at=datetime(2026, 8, 20, 12, 30, tzinfo=UTC),
        requester="Discord user 1001",
        filters="stato=all",
    )

    assert generated.filename.endswith(f".{export_format.value}")
    assert generated.record_count == 65
    assert generated.content
    if export_format is ExportFormat.PDF:
        reader = PdfReader(io.BytesIO(generated.content))
        assert len(reader.pages) >= 2
        assert "Fonte: Dipendenti in Cloud" in (reader.pages[0].extract_text() or "")
    elif export_format is ExportFormat.DOCX:
        document = Document(io.BytesIO(generated.content))
        assert document.tables
        assert len(document.tables[0].rows) == 66
        assert document.tables[0].cell(0, 0).text == "Nome"
        with zipfile.ZipFile(io.BytesIO(generated.content)) as archive:
            xml = archive.read("word/document.xml")
        assert b"tblHeader" in xml
    else:
        workbook = load_workbook(io.BytesIO(generated.content), data_only=False)
        try:
            sheet = workbook["Dipendenti"]
            assert sheet.freeze_panes == "A2"
            assert sheet.auto_filter.ref == "A1:G66"
            assert sheet["G2"].value == 1750
            assert sheet["G3"].value == "N/D"
            assert workbook["Riepilogo"]["B6"].value == "Dipendenti in Cloud"
        finally:
            workbook.close()


def test_ascii_table_is_complete_split_without_loss_and_marks_missing_values() -> None:
    rows = _rows(45)
    table = ascii_table(rows)
    chunks = split_ascii_for_discord(table, maximum=700)

    assert len(chunks) > 1
    assert "EMP-SYNTH-001" in table
    assert "EMP-SYNTH-045" in table
    assert "N/D" in table
    reconstructed = "\n".join(
        chunk.removeprefix("```text\n").removesuffix("\n```") for chunk in chunks
    )
    assert reconstructed == table


def test_export_size_limit_fails_closed() -> None:
    service = HrExportService(max_bytes=10)
    with pytest.raises(ExportGenerationError, match="size limit"):
        service.generate_employees(
            _rows(),
            export_format=ExportFormat.PDF,
            created_at=datetime(2026, 8, 20, tzinfo=UTC),
            requester="Discord user 1001",
            filters="stato=all",
        )


def test_xlsx_neutralizes_formula_like_identity_values() -> None:
    row = EmployeeExportRow(
        first_name='=HYPERLINK("https://invalid.example")',
        last_name="+cmd",
        employee_id="@EMP-SYNTH-001",
        status="active",
        contract_expiry=None,
        contract_type=None,
    )
    generated = HrExportService().generate_employees(
        (row,),
        export_format=ExportFormat.XLSX,
        created_at=datetime(2026, 8, 20, tzinfo=UTC),
        requester="Discord user 1001",
        filters="stato=all",
    )
    workbook = load_workbook(io.BytesIO(generated.content), data_only=False)
    try:
        sheet = workbook["Dipendenti"]
        assert str(sheet["A2"].value).startswith("'=")
        assert str(sheet["B2"].value).startswith("'+")
        assert str(sheet["C2"].value).startswith("'@")
        assert all(sheet.cell(2, column).data_type != "f" for column in range(1, 4))
    finally:
        workbook.close()
