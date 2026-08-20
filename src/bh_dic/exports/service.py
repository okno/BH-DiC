"""Deterministic PDF, DOCX, XLSX, and ASCII employee report generation."""

from __future__ import annotations

import io
import re
import zipfile
from collections.abc import Iterable, Sequence
from datetime import date, datetime
from decimal import Decimal
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook, load_workbook  # type: ignore[import-untyped]
from openpyxl.styles import Alignment, Font, PatternFill  # type: ignore[import-untyped]
from openpyxl.utils import get_column_letter  # type: ignore[import-untyped]
from reportlab.lib import colors  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import A4, landscape  # type: ignore[import-untyped]
from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.lib.units import cm  # type: ignore[import-untyped]
from reportlab.platypus import (  # type: ignore[import-untyped]
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from bh_dic.exports.models import EmployeeExportRow, ExportFormat, GeneratedExport

_HEADERS = (
    "Nome",
    "Cognome",
    "ID DiC",
    "Stato",
    "Scadenza contratto",
    "Tipologia contratto",
    "Netto mensile",
)
_SAFE_FILENAME = re.compile(r"[^a-z0-9._-]+")
_EXCEL_FORMULA_PREFIXES = ("=", "+", "-", "@")


class ExportGenerationError(RuntimeError):
    """A report could not be generated or validated locally."""


def _text(value: str | None) -> str:
    return value if value else "N/D"


def _date(value: date | None) -> str:
    if value is None:
        return "N/D"
    return value.strftime("%d/%m/%Y")


def _money(value: Decimal | None) -> str:
    if value is None:
        return "N/D"
    rendered = f"{value:,.2f}".replace(",", "_").replace(".", ",").replace("_", ".")
    return f"EUR {rendered}"


def _safe_excel_text(value: str) -> str:
    return f"'{value}" if value.startswith(_EXCEL_FORMULA_PREFIXES) else value


def _rows(rows: Iterable[EmployeeExportRow]) -> list[tuple[str, ...]]:
    return [
        (
            _text(row.first_name),
            _text(row.last_name),
            row.employee_id,
            row.status,
            _date(row.contract_expiry),
            _text(row.contract_type),
            _money(row.monthly_net),
        )
        for row in rows
    ]


def ascii_table(rows: Sequence[EmployeeExportRow]) -> str:
    """Return a lossless aligned table; missing values are rendered as ``N/D``."""

    body = _rows(rows)
    widths = [len(header) for header in _HEADERS]
    for values in body:
        for index, value in enumerate(values):
            widths[index] = max(widths[index], len(value))
    border = "+" + "+".join("-" * (width + 2) for width in widths) + "+"

    def line(values: Sequence[str]) -> str:
        return (
            "|"
            + "|".join(f" {value:<{width}} " for value, width in zip(values, widths, strict=True))
            + "|"
        )

    rendered = [border, line(_HEADERS), border]
    rendered.extend(line(values) for values in body)
    rendered.append(border)
    return "\n".join(rendered)


def split_ascii_for_discord(table: str, *, maximum: int = 1_900) -> tuple[str, ...]:
    """Split a table into numbered code blocks without dropping a physical line."""

    if maximum < 256:
        raise ValueError("Discord table chunk limit is too small")
    lines = table.splitlines()
    chunks: list[list[str]] = []
    current: list[str] = []
    current_size = 0
    for line in lines:
        if len(line) + 1 > maximum - 8:
            raise ExportGenerationError("one table row exceeds the Discord message limit")
        addition = len(line) + 1
        if current and current_size + addition > maximum - 8:
            chunks.append(current)
            current = []
            current_size = 0
        current.append(line)
        current_size += addition
    if current:
        chunks.append(current)
    return tuple(f"```text\n{'\n'.join(chunk)}\n```" for chunk in chunks)


class HrExportService:
    """Generate bounded reports in memory so no clear-text HR temp file reaches disk."""

    def __init__(self, *, max_bytes: int = 8 * 1024 * 1024) -> None:
        if max_bytes < 1:
            raise ValueError("export byte limit must be positive")
        self._max_bytes = max_bytes

    def generate_employees(
        self,
        rows: Sequence[EmployeeExportRow],
        *,
        export_format: ExportFormat,
        created_at: datetime,
        requester: str,
        filters: str,
        title: str = "Elenco dipendenti",
    ) -> GeneratedExport:
        safe_title = " ".join(title.split())[:160] or "Elenco dipendenti"
        safe_requester = " ".join(requester.split())[:128] or "N/D"
        safe_filters = " ".join(filters.split())[:512] or "Nessuno"
        content: bytes
        media_type: str
        if export_format is ExportFormat.PDF:
            content = self._pdf(rows, safe_title, created_at, safe_requester, safe_filters)
            media_type = "application/pdf"
        elif export_format is ExportFormat.DOCX:
            content = self._docx(rows, safe_title, created_at, safe_requester, safe_filters)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif export_format is ExportFormat.XLSX:
            content = self._xlsx(rows, safe_title, created_at, safe_requester, safe_filters)
            media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        else:  # pragma: no cover - closed enum defense
            raise ExportGenerationError("unsupported export format")
        if not content or len(content) > self._max_bytes:
            raise ExportGenerationError("generated export exceeds the configured size limit")
        timestamp = created_at.strftime("%Y-%m-%d_%H%M%S")
        stem = _SAFE_FILENAME.sub("_", safe_title.casefold()).strip("._-") or "dipendenti"
        filename = f"{stem[:80]}_{timestamp}.{export_format.value}"
        return GeneratedExport(filename, media_type, content, len(rows))

    @staticmethod
    def _metadata_lines(
        created_at: datetime,
        requester: str,
        filters: str,
        record_count: int,
    ) -> tuple[str, ...]:
        return (
            f"Generato: {created_at.strftime('%d/%m/%Y %H:%M:%S %Z')}",
            f"Richiedente: {requester}",
            f"Filtri: {filters}",
            f"Record: {record_count}",
            "Fonte: Dipendenti in Cloud",
        )

    @classmethod
    def _pdf(
        cls,
        rows: Sequence[EmployeeExportRow],
        title: str,
        created_at: datetime,
        requester: str,
        filters: str,
    ) -> bytes:
        output = io.BytesIO()
        document = SimpleDocTemplate(
            output,
            pagesize=landscape(A4),
            leftMargin=1 * cm,
            rightMargin=1 * cm,
            topMargin=1.2 * cm,
            bottomMargin=1.2 * cm,
            title=title,
        )
        styles = getSampleStyleSheet()
        story: list[Any] = [Paragraph(title, styles["Title"]), Spacer(1, 0.2 * cm)]
        for line in cls._metadata_lines(created_at, requester, filters, len(rows)):
            story.append(Paragraph(line, styles["BodyText"]))
        data = [list(_HEADERS), *[list(values) for values in _rows(rows)]]
        table = Table(
            data,
            repeatRows=1,
            colWidths=(2.6 * cm, 2.8 * cm, 2.3 * cm, 1.8 * cm, 3.1 * cm, 3.6 * cm, 2.5 * cm),
        )
        table.setStyle(
            TableStyle(
                (
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), (colors.white, colors.whitesmoke)),
                )
            )
        )
        story.extend((Spacer(1, 0.35 * cm), table))

        def footer(canvas: Any, doc: Any) -> None:
            canvas.saveState()
            canvas.setFont("Helvetica", 7)
            canvas.drawRightString(
                landscape(A4)[0] - 1 * cm,
                0.55 * cm,
                f"Pagina {doc.page}",
            )
            canvas.restoreState()

        document.build(story, onFirstPage=footer, onLaterPages=footer)
        content = output.getvalue()
        if not content.startswith(b"%PDF-"):
            raise ExportGenerationError("generated PDF failed local validation")
        return content

    @classmethod
    def _docx(
        cls,
        rows: Sequence[EmployeeExportRow],
        title: str,
        created_at: datetime,
        requester: str,
        filters: str,
    ) -> bytes:
        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = section.right_margin = Cm(1.2)
        document.add_heading(title, level=0)
        for line in cls._metadata_lines(created_at, requester, filters, len(rows)):
            paragraph = document.add_paragraph(line)
            if paragraph.style is not None:
                paragraph.style.font.size = Pt(9)
        table = document.add_table(rows=1, cols=len(_HEADERS))
        table.style = "Table Grid"
        header = table.rows[0]
        for cell, value in zip(header.cells, _HEADERS, strict=True):
            cell.text = value
            for run in cell.paragraphs[0].runs:
                run.bold = True
        properties = header._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        properties.append(repeat)
        for values in _rows(rows):
            cells = table.add_row().cells
            for cell, value in zip(cells, values, strict=True):
                cell.text = value
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7)
        output = io.BytesIO()
        document.save(output)
        content = output.getvalue()
        cls._validate_office_zip(content, required="word/document.xml")
        return content

    @classmethod
    def _xlsx(
        cls,
        rows: Sequence[EmployeeExportRow],
        title: str,
        created_at: datetime,
        requester: str,
        filters: str,
    ) -> bytes:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Dipendenti"
        sheet.append(_HEADERS)
        header_fill = PatternFill("solid", fgColor="D9EAF7")
        for cell in sheet[1]:
            cell.font = Font(bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
        for row in rows:
            sheet.append(
                (
                    _safe_excel_text(_text(row.first_name)),
                    _safe_excel_text(_text(row.last_name)),
                    _safe_excel_text(row.employee_id),
                    _safe_excel_text(row.status),
                    row.contract_expiry,
                    _safe_excel_text(_text(row.contract_type)),
                    row.monthly_net,
                )
            )
        for cell in sheet["E"][1:]:
            if cell.value is not None:
                cell.number_format = "dd/mm/yyyy"
        for cell in sheet["G"][1:]:
            if cell.value is None:
                cell.value = "N/D"
            else:
                cell.number_format = "[$EUR] #,##0.00"
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = f"A1:G{max(1, sheet.max_row)}"
        for index, header in enumerate(_HEADERS, start=1):
            values = [
                str(sheet.cell(row=row, column=index).value or "")
                for row in range(1, sheet.max_row + 1)
            ]
            width = min(45, max(len(header), *(len(value) for value in values)) + 2)
            sheet.column_dimensions[get_column_letter(index)].width = width
        summary = workbook.create_sheet("Riepilogo")
        summary.append(("Titolo", title))
        for line in cls._metadata_lines(created_at, requester, filters, len(rows)):
            key, _, value = line.partition(": ")
            summary.append((key, value))
        summary.column_dimensions["A"].width = 18
        summary.column_dimensions["B"].width = 70
        output = io.BytesIO()
        workbook.save(output)
        content = output.getvalue()
        cls._validate_office_zip(content, required="xl/workbook.xml")
        load_workbook(io.BytesIO(content), read_only=True, data_only=False).close()
        return content

    @staticmethod
    def _validate_office_zip(content: bytes, *, required: str) -> None:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = frozenset(archive.namelist())
        except (OSError, zipfile.BadZipFile) as exc:
            raise ExportGenerationError("generated Office document is invalid") from exc
        if required not in names or "[Content_Types].xml" not in names:
            raise ExportGenerationError("generated Office document is incomplete")
